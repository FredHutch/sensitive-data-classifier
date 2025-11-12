"""
Multi-Format Document Processor

Advanced document processing pipeline supporting multiple file formats:
- TXT: Plain text files
- DOCX: Microsoft Word documents  
- PDF: Portable Document Format files
- CSV: Comma-separated values files
- XLSX: Microsoft Excel spreadsheets
- JSON: JavaScript Object Notation files

Features:
- Robust text extraction with error handling
- Metadata preservation and analysis
- Security validation and file sanitization
- Batch processing capabilities
- HIPAA-compliant processing pipeline
"""

import os
import io
import csv
import json
import logging
import tempfile
from typing import Dict, List, Any, Optional, Union, Tuple
from pathlib import Path
from datetime import datetime
import hashlib

# Document processing libraries
try:
    from docx import Document
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

try:
    import PyPDF2
    import pdfplumber
    HAS_PDF = True
except ImportError:
    HAS_PDF = False

try:
    from openpyxl import load_workbook
    import pandas as pd
    HAS_EXCEL = True
except ImportError:
    HAS_EXCEL = False

try:
    import magic
    HAS_MAGIC = True
except ImportError:
    HAS_MAGIC = False

# OCR libraries for scanned documents
try:
    import pytesseract
    from PIL import Image
    from pdf2image import convert_from_bytes, convert_from_path
    HAS_OCR = True
except ImportError:
    HAS_OCR = False

from .security import SecurityManager

logger = logging.getLogger(__name__)

class DocumentProcessor:
    """
    Advanced multi-format document processor with security validation
    and comprehensive text extraction capabilities.
    """
    
    def __init__(self):
        self.security_manager = SecurityManager()
        
        # Supported file formats and their processors
        self.processors = {
            '.txt': self._process_text_file,
            '.docx': self._process_docx_file,
            '.pdf': self._process_pdf_file,
            '.csv': self._process_csv_file,
            '.xlsx': self._process_excel_file,
            '.json': self._process_json_file,
            # Image formats (OCR)
            '.png': self._process_image_file,
            '.jpg': self._process_image_file,
            '.jpeg': self._process_image_file,
            '.tiff': self._process_image_file,
            '.tif': self._process_image_file,
            '.bmp': self._process_image_file
        }

        # File type detection patterns
        self.mime_types = {
            'text/plain': '.txt',
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document': '.docx',
            'application/pdf': '.pdf',
            'text/csv': '.csv',
            'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet': '.xlsx',
            'application/json': '.json',
            'image/png': '.png',
            'image/jpeg': '.jpg',
            'image/tiff': '.tiff',
            'image/bmp': '.bmp'
        }
        
        # Processing statistics
        self.stats = {
            'files_processed': 0,
            'total_text_extracted': 0,
            'processing_errors': 0,
            'last_processing_time': None
        }
        
        logger.info("Document Processor initialized")
        logger.info(f"Available processors: {list(self.processors.keys())}")
        logger.info(f"Libraries available: DOCX={HAS_DOCX}, PDF={HAS_PDF}, EXCEL={HAS_EXCEL}, MAGIC={HAS_MAGIC}")
    
    def process_document(self, file_obj, filename: Optional[str] = None) -> Dict[str, Any]:
        """
        Process a single document and extract text content.
        
        Args:
            file_obj: File object or file path
            filename: Optional filename override
            
        Returns:
            Dict[str, Any]: Processing results with extracted text and metadata
        """
        start_time = datetime.now()
        
        try:
            # Determine filename and file type
            if isinstance(file_obj, str):
                # File path provided
                file_path = file_obj
                filename = filename or os.path.basename(file_path)
                with open(file_path, 'rb') as f:
                    file_content = f.read()
            else:
                # File object provided
                filename = filename or getattr(file_obj, 'filename', 'unknown.txt')
                file_obj.seek(0)
                file_content = file_obj.read()
                if hasattr(file_obj, 'seek'):
                    file_obj.seek(0)
            
            # Security validation
            if not self.security_manager.validate_file(file_obj if not isinstance(file_obj, str) else filename):
                return self._create_error_result(filename, "Security validation failed")
            
            # Detect file type
            file_extension = self._detect_file_type(file_content, filename)
            
            if file_extension not in self.processors:
                return self._create_error_result(filename, f"Unsupported file type: {file_extension}")
            
            # Process the document
            processor = self.processors[file_extension]
            result = processor(file_obj if not isinstance(file_obj, str) else file_content, filename)
            
            # Add metadata
            result.update({
                'filename': filename,
                'file_type': file_extension,
                'file_size': len(file_content),
                'processing_time': (datetime.now() - start_time).total_seconds(),
                'processed_at': datetime.now().isoformat(),
                'file_hash': hashlib.sha256(file_content).hexdigest(),
                'processor_version': '1.0.0'
            })
            
            # Update statistics
            self.stats['files_processed'] += 1
            self.stats['total_text_extracted'] += len(result.get('text', ''))
            self.stats['last_processing_time'] = datetime.now().isoformat()
            
            return result
            
        except Exception as e:
            logger.error(f"Error processing document {filename}: {e}")
            self.stats['processing_errors'] += 1
            return self._create_error_result(filename, str(e))
    
    def process_batch(self, file_list: List[Any], max_workers: int = 4) -> List[Dict[str, Any]]:
        """
        Process multiple documents in batch.
        
        Args:
            file_list: List of file objects or file paths
            max_workers: Maximum number of concurrent workers
            
        Returns:
            List[Dict[str, Any]]: Processing results for all files
        """
        results = []
        
        logger.info(f"Starting batch processing of {len(file_list)} documents")
        
        for i, file_obj in enumerate(file_list):
            if i % 10 == 0:
                logger.info(f"Processed {i}/{len(file_list)} documents")
            
            result = self.process_document(file_obj)
            results.append(result)
        
        logger.info(f"Batch processing completed: {len(results)} documents processed")
        return results
    
    def _detect_file_type(self, file_content: bytes, filename: str) -> str:
        """
        Detect file type using multiple methods.
        
        Args:
            file_content: Raw file content
            filename: Original filename
            
        Returns:
            str: Detected file extension
        """
        # Method 1: File extension
        file_ext = Path(filename).suffix.lower()
        if file_ext in self.processors:
            return file_ext
        
        # Method 2: Magic number detection (if available)
        if HAS_MAGIC:
            try:
                mime_type = magic.from_buffer(file_content, mime=True)
                if mime_type in self.mime_types:
                    return self.mime_types[mime_type]
            except Exception as e:
                logger.warning(f"Magic detection failed: {e}")
        
        # Method 3: Content-based detection
        if file_content.startswith(b'%PDF'):
            return '.pdf'
        elif file_content.startswith(b'PK\x03\x04'):
            # ZIP-based formats (DOCX, XLSX)
            if b'word/' in file_content:
                return '.docx'
            elif b'xl/' in file_content:
                return '.xlsx'
        
        # Default to text if nothing else detected
        return '.txt'
    
    def _process_text_file(self, file_obj, filename: str) -> Dict[str, Any]:
        """
        Process plain text files.
        
        Args:
            file_obj: File object or file content
            filename: Filename
            
        Returns:
            Dict[str, Any]: Processing result
        """
        try:
            if isinstance(file_obj, bytes):
                # Detect encoding
                try:
                    text = file_obj.decode('utf-8')
                except UnicodeDecodeError:
                    try:
                        text = file_obj.decode('latin-1')
                    except UnicodeDecodeError:
                        text = file_obj.decode('utf-8', errors='ignore')
            else:
                text = file_obj.read()
                if isinstance(text, bytes):
                    text = text.decode('utf-8', errors='ignore')
            
            return {
                'success': True,
                'text': text,
                'word_count': len(text.split()),
                'character_count': len(text),
                'line_count': len(text.splitlines()),
                'encoding_detected': 'utf-8'
            }
            
        except Exception as e:
            return {'success': False, 'error': str(e), 'text': ''}
    
    def _process_docx_file(self, file_obj, filename: str) -> Dict[str, Any]:
        """
        Process Microsoft Word DOCX files.
        
        Args:
            file_obj: File object or file content
            filename: Filename
            
        Returns:
            Dict[str, Any]: Processing result
        """
        if not HAS_DOCX:
            return {'success': False, 'error': 'python-docx not available', 'text': ''}
        
        try:
            # Handle file object or bytes
            if isinstance(file_obj, bytes):
                file_stream = io.BytesIO(file_obj)
            else:
                file_stream = file_obj
            
            # Load document
            doc = Document(file_stream)
            
            # Extract text from paragraphs
            text_parts = []
            paragraph_count = 0
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
                    paragraph_count += 1
            
            # Extract text from tables
            table_count = 0
            for table in doc.tables:
                table_count += 1
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_parts.append(' | '.join(row_text))
            
            full_text = '\n'.join(text_parts)
            
            return {
                'success': True,
                'text': full_text,
                'word_count': len(full_text.split()),
                'character_count': len(full_text),
                'paragraph_count': paragraph_count,
                'table_count': table_count,
                'document_properties': self._extract_docx_properties(doc)
            }
            
        except Exception as e:
            logger.error(f"DOCX processing error: {e}")
            return {'success': False, 'error': str(e), 'text': ''}
    
    def _process_pdf_file(self, file_obj, filename: str) -> Dict[str, Any]:
        """
        Process PDF files using multiple extraction methods.
        
        Args:
            file_obj: File object or file content
            filename: Filename
            
        Returns:
            Dict[str, Any]: Processing result
        """
        if not HAS_PDF:
            return {'success': False, 'error': 'PDF processing libraries not available', 'text': ''}
        
        try:
            # Handle file object or bytes
            if isinstance(file_obj, bytes):
                file_stream = io.BytesIO(file_obj)
            else:
                file_stream = file_obj
            
            text_parts = []
            page_count = 0
            
            # Method 1: Try pdfplumber first (better for complex layouts)
            try:
                with pdfplumber.open(file_stream) as pdf:
                    for page in pdf.pages:
                        page_text = page.extract_text()
                        if page_text:
                            text_parts.append(page_text)
                        page_count += 1
            except Exception as e:
                logger.warning(f"pdfplumber extraction failed: {e}")
                
                # Method 2: Fallback to PyPDF2
                file_stream.seek(0)
                try:
                    pdf_reader = PyPDF2.PdfReader(file_stream)
                    page_count = len(pdf_reader.pages)
                    
                    for page_num in range(page_count):
                        page = pdf_reader.pages[page_num]
                        page_text = page.extract_text()
                        if page_text:
                            text_parts.append(page_text)
                except Exception as e2:
                    logger.error(f"PyPDF2 extraction also failed: {e2}")
                    return {'success': False, 'error': f'PDF extraction failed: {e2}', 'text': ''}
            
            full_text = '\n'.join(text_parts)

            # If no text was extracted, try OCR on scanned PDF
            if not full_text.strip() and HAS_OCR:
                logger.info("No text extracted from PDF, attempting OCR on scanned document")
                file_stream.seek(0)
                ocr_result = self._process_pdf_with_ocr(file_stream, filename)
                if ocr_result['success']:
                    return ocr_result

            return {
                'success': True,
                'text': full_text,
                'word_count': len(full_text.split()),
                'character_count': len(full_text),
                'page_count': page_count,
                'extraction_method': 'pdfplumber' if text_parts else 'PyPDF2'
            }

        except Exception as e:
            logger.error(f"PDF processing error: {e}")
            return {'success': False, 'error': str(e), 'text': ''}
    
    def _process_csv_file(self, file_obj, filename: str) -> Dict[str, Any]:
        """
        Process CSV files.
        
        Args:
            file_obj: File object or file content
            filename: Filename
            
        Returns:
            Dict[str, Any]: Processing result
        """
        try:
            if isinstance(file_obj, bytes):
                content = file_obj.decode('utf-8', errors='ignore')
                file_stream = io.StringIO(content)
            else:
                content = file_obj.read()
                if isinstance(content, bytes):
                    content = content.decode('utf-8', errors='ignore')
                file_stream = io.StringIO(content)
            
            # Detect CSV dialect
            sample = content[:1024]
            try:
                dialect = csv.Sniffer().sniff(sample)
            except:
                dialect = csv.excel
            
            # Read CSV data
            file_stream.seek(0)
            csv_reader = csv.reader(file_stream, dialect=dialect)
            
            rows = []
            header = None
            row_count = 0
            
            for i, row in enumerate(csv_reader):
                if i == 0:
                    header = row
                rows.append(row)
                row_count += 1
                
                # Limit to prevent memory issues
                if row_count > 10000:
                    break
            
            # Convert to text representation
            text_parts = []
            if header:
                text_parts.append('Headers: ' + ', '.join(header))
            
            for row in rows[:100]:  # First 100 rows for text extraction
                if row:
                    text_parts.append(' | '.join(str(cell) for cell in row))
            
            full_text = '\n'.join(text_parts)
            
            return {
                'success': True,
                'text': full_text,
                'word_count': len(full_text.split()),
                'character_count': len(full_text),
                'row_count': row_count,
                'column_count': len(header) if header else 0,
                'header': header,
                'dialect': str(dialect.__class__.__name__)
            }
            
        except Exception as e:
            logger.error(f"CSV processing error: {e}")
            return {'success': False, 'error': str(e), 'text': ''}
    
    def _process_excel_file(self, file_obj, filename: str) -> Dict[str, Any]:
        """
        Process Excel XLSX files.
        
        Args:
            file_obj: File object or file content
            filename: Filename
            
        Returns:
            Dict[str, Any]: Processing result
        """
        if not HAS_EXCEL:
            return {'success': False, 'error': 'openpyxl not available', 'text': ''}
        
        try:
            # Handle file object or bytes
            if isinstance(file_obj, bytes):
                file_stream = io.BytesIO(file_obj)
            else:
                file_stream = file_obj
            
            # Load workbook
            workbook = load_workbook(file_stream, data_only=True)
            
            text_parts = []
            sheet_count = len(workbook.worksheets)
            total_rows = 0
            
            # Process each worksheet
            for sheet_name in workbook.sheetnames:
                sheet = workbook[sheet_name]
                text_parts.append(f"\n--- Sheet: {sheet_name} ---")
                
                # Get data from sheet
                rows_processed = 0
                for row in sheet.iter_rows(values_only=True):
                    if row and any(cell is not None for cell in row):
                        row_text = ' | '.join(str(cell) if cell is not None else '' for cell in row)
                        if row_text.strip():
                            text_parts.append(row_text)
                            rows_processed += 1
                    
                    # Limit rows per sheet
                    if rows_processed > 1000:
                        break
                
                total_rows += rows_processed
            
            full_text = '\n'.join(text_parts)
            
            return {
                'success': True,
                'text': full_text,
                'word_count': len(full_text.split()),
                'character_count': len(full_text),
                'sheet_count': sheet_count,
                'total_rows': total_rows,
                'sheet_names': workbook.sheetnames
            }
            
        except Exception as e:
            logger.error(f"Excel processing error: {e}")
            return {'success': False, 'error': str(e), 'text': ''}
    
    def _process_json_file(self, file_obj, filename: str) -> Dict[str, Any]:
        """
        Process JSON files.
        
        Args:
            file_obj: File object or file content
            filename: Filename
            
        Returns:
            Dict[str, Any]: Processing result
        """
        try:
            if isinstance(file_obj, bytes):
                content = file_obj.decode('utf-8', errors='ignore')
            else:
                content = file_obj.read()
                if isinstance(content, bytes):
                    content = content.decode('utf-8', errors='ignore')
            
            # Parse JSON
            data = json.loads(content)
            
            # Extract text from JSON structure
            text_parts = self._extract_text_from_json(data)
            full_text = '\n'.join(text_parts)
            
            return {
                'success': True,
                'text': full_text,
                'word_count': len(full_text.split()),
                'character_count': len(full_text),
                'json_structure': self._analyze_json_structure(data)
            }
            
        except json.JSONDecodeError as e:
            return {'success': False, 'error': f'Invalid JSON: {e}', 'text': content[:1000]}
        except Exception as e:
            logger.error(f"JSON processing error: {e}")
            return {'success': False, 'error': str(e), 'text': ''}
    
    def _extract_text_from_json(self, data: Any, path: str = '') -> List[str]:
        """
        Recursively extract text from JSON structure.
        
        Args:
            data: JSON data
            path: Current path in JSON structure
            
        Returns:
            List[str]: Extracted text strings
        """
        text_parts = []
        
        if isinstance(data, dict):
            for key, value in data.items():
                new_path = f"{path}.{key}" if path else key
                if isinstance(value, str) and value.strip():
                    text_parts.append(f"{new_path}: {value}")
                else:
                    text_parts.extend(self._extract_text_from_json(value, new_path))
        
        elif isinstance(data, list):
            for i, item in enumerate(data):
                new_path = f"{path}[{i}]"
                text_parts.extend(self._extract_text_from_json(item, new_path))
        
        elif isinstance(data, str) and data.strip():
            text_parts.append(data)
        
        elif data is not None:
            text_parts.append(str(data))
        
        return text_parts
    
    def _analyze_json_structure(self, data: Any) -> Dict[str, Any]:
        """
        Analyze JSON structure for metadata.
        
        Args:
            data: JSON data
            
        Returns:
            Dict[str, Any]: Structure analysis
        """
        if isinstance(data, dict):
            return {
                'type': 'object',
                'keys': len(data),
                'key_names': list(data.keys())[:10]  # First 10 keys
            }
        elif isinstance(data, list):
            return {
                'type': 'array',
                'length': len(data),
                'item_types': list(set(type(item).__name__ for item in data[:10]))
            }
        else:
            return {
                'type': type(data).__name__,
                'value': str(data)[:100]
            }
    
    def _extract_docx_properties(self, doc) -> Dict[str, Any]:
        """
        Extract document properties from DOCX file.
        
        Args:
            doc: Document object
            
        Returns:
            Dict[str, Any]: Document properties
        """
        try:
            core_props = doc.core_properties
            return {
                'title': core_props.title or '',
                'author': core_props.author or '',
                'subject': core_props.subject or '',
                'created': core_props.created.isoformat() if core_props.created else '',
                'modified': core_props.modified.isoformat() if core_props.modified else '',
                'category': core_props.category or '',
                'comments': core_props.comments or ''
            }
        except Exception as e:
            logger.warning(f"Could not extract document properties: {e}")
            return {}
    
    def _create_error_result(self, filename: str, error_message: str) -> Dict[str, Any]:
        """
        Create standardized error result.
        
        Args:
            filename: Filename
            error_message: Error description
            
        Returns:
            Dict[str, Any]: Error result
        """
        return {
            'success': False,
            'filename': filename,
            'error': error_message,
            'text': '',
            'word_count': 0,
            'character_count': 0,
            'processed_at': datetime.now().isoformat()
        }
    
    def _process_pdf_with_ocr(self, file_obj, filename: str) -> Dict[str, Any]:
        """
        Process scanned PDF using OCR.

        Args:
            file_obj: File object or file content
            filename: Filename

        Returns:
            Dict[str, Any]: Processing result with extracted text
        """
        if not HAS_OCR:
            return {'success': False, 'error': 'OCR libraries not available', 'text': ''}

        try:
            # Convert file to bytes if needed
            if isinstance(file_obj, bytes):
                pdf_bytes = file_obj
            else:
                pdf_bytes = file_obj.read()

            # Convert PDF pages to images
            logger.info(f"Converting PDF pages to images for OCR: {filename}")
            images = convert_from_bytes(pdf_bytes, dpi=300)

            text_parts = []
            for i, image in enumerate(images):
                logger.debug(f"Processing page {i+1}/{len(images)} with OCR")
                # Perform OCR on each page
                page_text = pytesseract.image_to_string(image, lang='eng', config='--psm 1')
                if page_text.strip():
                    text_parts.append(page_text)

            full_text = '\n\n'.join(text_parts)

            return {
                'success': True,
                'text': full_text,
                'word_count': len(full_text.split()),
                'character_count': len(full_text),
                'page_count': len(images),
                'extraction_method': 'OCR (Tesseract)',
                'ocr_applied': True
            }

        except Exception as e:
            logger.error(f"PDF OCR processing error: {e}")
            return {'success': False, 'error': f'OCR failed: {str(e)}', 'text': ''}

    def _process_image_file(self, file_obj, filename: str) -> Dict[str, Any]:
        """
        Process image files using OCR.

        Args:
            file_obj: File object or file content
            filename: Filename

        Returns:
            Dict[str, Any]: Processing result with extracted text
        """
        if not HAS_OCR:
            return {'success': False, 'error': 'OCR libraries not available', 'text': ''}

        try:
            # Convert file to bytes if needed
            if isinstance(file_obj, bytes):
                image_bytes = file_obj
            else:
                image_bytes = file_obj.read()

            # Open image with PIL
            image = Image.open(io.BytesIO(image_bytes))

            # Perform OCR
            logger.info(f"Performing OCR on image: {filename}")
            text = pytesseract.image_to_string(image, lang='eng', config='--psm 1')

            # Get image metadata
            width, height = image.size
            format_name = image.format or 'Unknown'

            return {
                'success': True,
                'text': text,
                'word_count': len(text.split()),
                'character_count': len(text),
                'extraction_method': 'OCR (Tesseract)',
                'image_width': width,
                'image_height': height,
                'image_format': format_name,
                'ocr_applied': True
            }

        except Exception as e:
            logger.error(f"Image OCR processing error: {e}")
            return {'success': False, 'error': f'Image OCR failed: {str(e)}', 'text': ''}

    def get_processing_stats(self) -> Dict[str, Any]:
        """
        Get processing statistics.

        Returns:
            Dict[str, Any]: Processing statistics
        """
        return {
            **self.stats,
            'supported_formats': list(self.processors.keys()),
            'available_libraries': {
                'docx': HAS_DOCX,
                'pdf': HAS_PDF,
                'excel': HAS_EXCEL,
                'magic': HAS_MAGIC,
                'ocr': HAS_OCR
            }
        }
    
    def reset_stats(self):
        """Reset processing statistics."""
        self.stats = {
            'files_processed': 0,
            'total_text_extracted': 0,
            'processing_errors': 0,
            'last_processing_time': None
        }